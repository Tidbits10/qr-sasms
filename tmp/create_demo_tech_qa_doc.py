from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\qr-sasms\QR-SASMS_Demo_Tech_Stack_and_Functionality_QA.docx"
NAVY, BLUE, LIGHT, MUTED = "0B2545", "2E74B5", "E8EEF5", "5B6573"
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.85); sec.bottom_margin = Inches(.75); sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)

styles=doc.styles
base=styles['Normal']; base.font.name='Calibri'; base._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); base._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); base.font.size=Pt(10.5); base.paragraph_format.space_after=Pt(5); base.paragraph_format.line_spacing=1.1
for n,s,b,a,c in [('Heading 1',16,15,7,BLUE),('Heading 2',13,10,5,BLUE),('Heading 3',11.5,7,3,NAVY)]:
    st=styles[n]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(s); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(c); st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a)

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd'); e.set(qn('w:fill'),fill); pr.append(e)
def callout(title,text):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; cell=t.cell(0,0); shade(cell,LIGHT)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(title+' '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def qa(question, answer):
    doc.add_heading(question,level=2)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); r=p.add_run('Suggested answer: '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run('“'+answer+'”')

hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=hp.add_run('QR-SASMS | Demo Technical Q&A'); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run('QR-SASMS Capstone Research Defense | Page '); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(75); p.paragraph_format.space_after=Pt(8); r=p.add_run('QR-SASMS'); r.font.size=Pt(28); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20); r=p.add_run('Demo Technical Q&A Reviewer\nTech Stack and System Functionalities'); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
callout('How to answer:', 'Start with the process in simple terms, then mention the technology only when needed. Keep the answer direct and avoid claiming that a feature is fully production-ready unless it has been tested in production.')
doc.add_page_break()

doc.add_heading('General Answer Structure',level=1)
doc.add_paragraph('Use this line when the panel asks a technical question after the demo:')
callout('Opening line:', 'The feature shown in the demo is supported by our frontend, backend, database, and security components.')

doc.add_heading('Technical Stack',level=1)
qa('What tech stack did you use?','We used Next.js 14 as our full-stack web framework. It uses React for the frontend interface and Node.js through API routes for backend processing. We used TypeScript and JavaScript for application logic, PostgreSQL as the relational database, and Prisma ORM to connect and manage database records. We also used bcrypt for password hashing, JWT or jose for secure authentication and QR token validation, CSS for the interface design, GitHub for version control, and Render for deployment.')

doc.add_heading('Request and Status Processing',level=1)
qa('How does request submission work technically?','When a student submits a request, the frontend sends the form data to a protected API route. The backend validates the data and saves the request in PostgreSQL through Prisma. The request is initially marked as Pending, then it becomes visible in the Admin Staff processing queue.')
qa('How does the status update work?','The Admin Staff updates the request through authorized API routes. The backend checks the staff role, updates the request status in the database, creates a notification for the student, and records important actions in the audit log.')
qa('What happens when a request is rejected?','The Admin Staff enters a rejection reason. The student sees the status and reason in the dashboard, then can correct or re-upload the requirements and submit again.')

doc.add_heading('QR Verification',level=1)
qa('How does QR verification work?','When a request becomes Ready to Claim, the backend creates a unique secure QR token. The QR code does not directly expose the student’s personal information. During claiming, Scanner Personnel scan the QR code, the scanner sends the token to the server, and the server checks if it is valid, unexpired, and unused. If valid, the system marks the request as Completed and logs the transaction.')
qa('Why use QR codes for document claiming?','QR verification makes claiming faster and traceable. The server validates a one-time secure token before the document is released, which reduces unauthorized or repeated claims.')

doc.add_heading('Database and Prisma',level=1)
qa('Why use PostgreSQL?','We selected PostgreSQL because QR-SASMS handles related records such as users, requests, appointments, organizations, notifications, and audit logs. A relational database helps maintain consistency between those records.')
qa('What is Prisma?','Prisma is our ORM, or Object Relational Mapper. It allows our Next.js backend to communicate with PostgreSQL through structured code. It also helps us manage database models, relationships, validation, and queries.')

doc.add_heading('Appointment Controls',level=1)
qa('How do you prevent duplicate appointments?','Before an appointment is saved, the backend checks the date, time slot, configured capacity, and existing appointments. If the slot is already full or the student already has an appointment on that date, the booking is rejected. The unavailable slot is shown as disabled in the interface.')

doc.add_heading('Security and Roles',level=1)
qa('What are your security features?','The system uses hashed passwords through bcrypt, authenticated sessions, role-based authorization, secure server-validated QR tokens, audit logs, restricted staff functions, and controlled access to student requests and uploaded files.')
qa('Why role-based access control?','Role-based access ensures that each user only sees the functions needed for their responsibilities. Students submit and track requests, Admin Staff process operations, Super Admin manages sensitive configurations and accounts, and Scanner Personnel only verify QR claims.')
qa('Why does Event Request have restricted access?','Event Requests represent official student organizations, so not every student can submit them. The Super Admin assigns verified officers as Organization Representatives. Only active assigned representatives can access the Event Request function.')

doc.add_heading('Chatbot and Notifications',level=1)
qa('What happens if the chatbot cannot answer?','The chatbot provides FAQ-based answers first. If the concern is unique or cannot be resolved through the FAQ, it can be escalated into a Help Desk ticket where Admin Staff can provide a formal response.')
qa('How do notifications work?','When a major transaction happens, such as request submission, approval, rejection, Ready to Claim status, appointment update, or Help Desk response, the backend creates a notification record. The user sees it through the notification bell and dashboard.')

doc.add_heading('Deployment and Email',level=1)
qa('Does Render send the email?','Render hosts the frontend and backend application. Email delivery is handled by a separate SMTP email provider configured through environment variables. This separation keeps email credentials outside the source code.')
qa('What are the limitations?','For full institutional production use, the system still requires official university deployment approval, stable production database hosting, verified email configuration, controlled user acceptance testing, and possible integration with the official student information system.')

doc.add_heading('All-in-One Technical Answer',level=1)
callout('Memorize this response:', 'Our frontend collects and displays data, our protected backend API routes validate and process transactions, PostgreSQL stores the records through Prisma, and role-based access plus QR token validation protect sensitive actions. This allows the system to support the complete workflow from submission to secure claiming.')

doc.core_properties.title='QR-SASMS Demo Tech Stack and Functionality Q&A'; doc.core_properties.author='QR-SASMS Capstone Team'; doc.save(OUT); print(OUT)
