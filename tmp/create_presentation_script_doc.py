from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\qr-sasms\QR-SASMS_Presentation_Script_and_Reviewer_Guide.docx"
NAVY, BLUE, LIGHT, MUTED = "0B2545", "2E74B5", "E8EEF5", "5B6573"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.85); sec.bottom_margin = Inches(.75)
sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.1
for name, size, before, after, color in [('Heading 1',16,15,7,BLUE),('Heading 2',13,10,5,BLUE),('Heading 3',11.5,7,3,NAVY)]:
    st=styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)

def shade(cell, color):
    pr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), color); pr.append(shd)
def margins(cell):
    pr=cell._tc.get_or_add_tcPr(); mar=OxmlElement('w:tcMar')
    for side in ['top','start','bottom','end']:
        e=OxmlElement('w:'+side); e.set(qn('w:w'),'100' if side in ['top','bottom'] else '120'); e.set(qn('w:type'),'dxa'); mar.append(e)
    pr.append(mar)
def table_style(table):
    table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; table.style='Table Grid'
    for row in table.rows:
        for cell in row.cells: cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
def callout(title,text):
    t=doc.add_table(rows=1,cols=1); table_style(t); c=t.cell(0,0); shade(c,LIGHT)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(title+' '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def add_script(heading, text):
    doc.add_heading(heading, level=2)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    p.add_run('Script: ').bold=True
    p.add_run('“'+text+'”')

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('QR-SASMS | Presentation Script'); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('QR-SASMS Capstone Research Defense | Page '); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED); f=OxmlElement('w:fldSimple'); f.set(qn('w:instr'),'PAGE'); footer._p.append(f)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(74); p.paragraph_format.space_after=Pt(8)
r=p.add_run('QR-SASMS'); r.font.size=Pt(28); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20); r=p.add_run('Capstone Research Defense Presentation Script\nand Developer Reviewer Placement Guide'); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
callout('How to use this guide:', 'Deliver the scripts naturally in your own voice. Do not memorize every word. The reviewer topics at the end are mainly for the panel question-and-answer portion.')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50); r=p.add_run('QR-Based Student Affairs and Services Management System\nPUP San Pedro Campus'); r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(MUTED)
doc.add_page_break()

doc.add_heading('Slide-by-Slide Presentation Script', level=1)
scripts=[
('Slide 1 - Title','Good morning, panelists. We are presenting our proposed system entitled QR-SASMS, or the QR-Based Student Affairs and Services Management System for PUP San Pedro Campus. Our system aims to improve how student services are requested, processed, monitored, and claimed through a centralized web-based platform with QR code verification.'),
('Slide 2 - Background of the Study','Currently, student-service transactions may involve manual forms, repeated follow-ups, long queues, and delayed updates. Students may not immediately know if their request is pending, approved, rejected, or ready to claim. Because of this, we proposed QR-SASMS. The system allows students to submit requests online, book appointments, track their request status, receive notifications, and securely claim documents using QR verification.'),
('Slide 3 - Problems / Research Questions','Our study addresses problems such as inefficient manual processing, lack of real-time request tracking, duplicate appointment booking, delayed communication, and difficulty verifying document release. The system answers these problems by providing request tracking, appointment capacity control, notifications, role-based access, and QR-based claim verification.'),
('Slide 4 - Objectives','The general objective of our study is to develop a centralized Student Affairs and Services Management System for PUP San Pedro Campus. Specifically, the system aims to allow students to submit and monitor requests online, allow staff to manage transactions efficiently, prevent duplicate appointments, and provide a secure QR verification process for document claiming.'),
('Slide 5 - Significance of the Study','For students, QR-SASMS provides convenience because they can submit requests and track updates without repeatedly visiting the office. For Student Services personnel, it provides a centralized processing queue, request monitoring, appointment management, and audit logs. For the institution, it supports a more organized, transparent, and traceable student-service workflow.'),
('Slide 6 - Scope and Limitations','The system covers student requests, appointment booking, ID applications, help desk concerns, complaints, referrals, event requests for authorized organization representatives, notifications, QR claiming, reporting, and administrative management. The Event Request is not available to every student. Only a student assigned by the Super Admin as an active Organization Representative can access and submit an Event Request. For real deployment, the system still requires proper production configuration, a stable database, secure environment variables, and an SMTP provider for real email notifications.'),
('Slide 7 - Participants / Data Gathering','Our target users include students, Student Services personnel, Admin Staff, Super Admin, and Scanner Personnel. The system uses role-based access so that every user can only access the functions related to their responsibilities.'),
('Slide 8 - System Architecture Diagram','This is the system architecture of QR-SASMS. At the user interface layer, we have the Student Portal, Admin Staff Dashboard, Super Admin Console, and QR Scanner page. These interfaces communicate with the application layer, which is built using Next.js, React, Node.js, and API routes. The application layer handles authentication, request processing, appointment management, notifications, and QR validation. For the data layer, we use PostgreSQL as our relational database and Prisma ORM to manage database transactions and relationships.'),
('Slide 9 - Use Case Diagram','This Use Case Diagram shows the interaction between the different users and the system. The Student can submit and track requests, book appointments, view notifications, send Help Desk or Complaint concerns, and submit Event Requests only if assigned as an Organization Representative. Admin Staff processes requests, manages appointments, and handles tickets. The Super Admin manages staff accounts, masterlist records, organizations, system settings, reports, and backups. Scanner Personnel have limited access. Their role is to verify QR codes and complete document claims.'),
('Slide 10 - System Flowchart','This flowchart explains the end-to-end transaction process. First, the student registers or logs in. The student then selects a service, submits a request, or books an appointment. The request enters the Admin Staff processing queue. Staff review the request and either approve or reject it. If rejected, the student can see the reason and correct or re-upload the required documents. Once approved and ready to claim, the system generates a secure QR token. At the office, Scanner Personnel scan the QR code. The server validates it, then marks the request as completed and records the action in the audit log.'),
('Slide 11 - Expected Output / Features','The expected output is a functional web-based platform where students can manage requests and appointments while staff can process transactions in one centralized system. Its key features include request tracking, appointment booking, QR claiming, role-based access, notifications, audit logs, reports, FAQ chatbot, Help Desk escalation, and organization representative authorization.'),
('Slide 12 - Conclusion','In conclusion, QR-SASMS is designed to reduce manual processing and improve the efficiency, transparency, and security of student services. It provides students with better access to services while helping the Student Services Office organize, monitor, and verify transactions more effectively.'),
('Slide 13 - Closing','Thank you, panelists. We are now ready for your questions and recommendations.'),
]
for h,t in scripts: add_script(h,t)

doc.add_page_break()
doc.add_heading('Where to Use the Developer Reviewer Content',level=1)
doc.add_paragraph('Do not present the entire developer reviewer in the main slide presentation. Use the following map so technical details are stated only where they are relevant.')
table=doc.add_table(rows=1,cols=2); table_style(table)
for i,head in enumerate(['Reviewer topic','Where to mention it']): shade(table.rows[0].cells[i],LIGHT); table.rows[0].cells[i].paragraphs[0].add_run(head).bold=True
rows=[
('System overview','Slides 1 and 2'),('Problems solved','Slide 3'),('Main objectives','Slide 4'),('Benefits and significance','Slide 5'),('Roles and access control','Slides 6 and 9'),('Technology stack','Slide 8'),('PostgreSQL and Prisma','Slide 8'),('QR security and one-time validation','Slides 9 and 10'),('Appointment double-booking prevention','Slide 3 or 10; explain further during Q&A'),('Organization Representatives','Slides 6 and 9'),('Chatbot-to-Help Desk ticket escalation','Slide 11'),('Email, SMTP, Render, and deployment','Q&A unless your PPT has a deployment slide'),('bcrypt, JWT, authorization, and audit logs','Q&A'),('Database relationships and limitations','Q&A or Slide 6'),
]
for a,b in rows:
    c=table.add_row().cells; c[0].text=a; c[1].text=b

doc.add_heading('Developer Q&A Reminders',level=1)
for item in [
    'Explain the process first before giving the technical term.',
    'For QR questions: state that the server validates a unique token, not visible personal data.',
    'For email questions: explain that Render hosts the system; an external SMTP provider sends emails.',
    'For organization Event Requests: emphasize that only active assigned representatives can access it.',
    'For security questions: mention bcrypt password hashing, role-based authorization, QR validation, and audit logs.',
    'Do not claim that a feature is unlimited, fully automatic, or production-ready unless it has been tested and configured in production.',
]:
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(item)

callout('Presentation tip:', 'Keep the main presentation clear and focused on the problem, solution, architecture, use cases, and flowchart. Reserve deeper implementation details for the panel questions.')

doc.core_properties.title='QR-SASMS Presentation Script and Reviewer Guide'; doc.core_properties.subject='Capstone defense presentation script'; doc.core_properties.author='QR-SASMS Capstone Team'
doc.save(OUT)
print(OUT)
