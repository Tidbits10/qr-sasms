from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\qr-sasms\QR-SASMS_Developer_Defense_Reviewer.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
GOLD = "C99500"
MUTED = "5B6573"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 15, 7),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 11.5, NAVY, 7, 3),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")

def clean_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def add_label_paragraph(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    return p

def add_callout(title, text):
    t = doc.add_table(rows=1, cols=1)
    clean_table(t)
    cell = t.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_qa(question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Q: " + question)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Suggested answer: ")
    r2.bold = True
    p2.add_run(answer)

# Header and footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("QR-SASMS | Developer Defense Reviewer")
hr.font.size = Pt(8.5)
hr.font.color.rgb = RGBColor.from_string(MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("QR-SASMS Capstone Research Defense | Page ")
fr.font.size = Pt(8.5)
fr.font.color.rgb = RGBColor.from_string(MUTED)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(78)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("QR-SASMS")
r.bold = True
r.font.name = "Calibri"
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("Developer Reviewer for Capstone Research Defense")
r.font.size = Pt(17)
r.font.color.rgb = RGBColor.from_string(BLUE)

add_callout("Purpose:", "Use this guide to explain the system design, technical stack, workflow, security controls, and likely panel questions during the proposal or final defense.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run("QR-Based Student Affairs and Services Management System\nPUP San Pedro Campus")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(MUTED)

doc.add_page_break()

doc.add_heading("1. System Overview", level=1)
doc.add_paragraph(
    "QR-SASMS stands for QR-Based Student Affairs and Services Management System. It is a web-based platform designed for the PUP San Pedro Student Services Office. The system centralizes common student-service transactions so students can submit requests, book appointments, track progress, receive notifications, and claim approved documents through secure QR verification."
)
doc.add_paragraph("The system addresses common manual-process issues:")
for item in [
    "Manual forms, repeated follow-ups, and long queues.",
    "Delayed request status updates.",
    "Duplicate appointment booking.",
    "Unverified or undocumented document release.",
    "Difficulty monitoring complaints, referrals, help desk concerns, and organization event requests.",
]: add_bullet(item)

doc.add_heading("2. Core System Flow", level=1)
doc.add_paragraph("Memorize this end-to-end flow:")
for item in [
    "Student registers or logs in.",
    "Student submits a service request or books an appointment.",
    "The system validates the submission and stores it in PostgreSQL.",
    "Admin Staff reviews the request through the processing queue.",
    "The request status moves through Pending, Approved, Ready to Claim, Rejected, or Completed.",
    "For a ready document, the system generates a secure QR token.",
    "Scanner Personnel verifies the QR token at the office.",
    "A valid one-time claim is completed and recorded in the audit log.",
]: add_number(item)

add_callout("Rejected-request flow:", "Pending → Rejected with reason → Student corrects or re-uploads requirements → Resubmits for staff review.")

doc.add_heading("3. User Roles and Access", level=1)
table = doc.add_table(rows=1, cols=2)
clean_table(table)
table.style = "Table Grid"
for i, head in enumerate(["Role", "Main access"]):
    cell = table.rows[0].cells[i]
    shade(cell, LIGHT_BLUE)
    cell.paragraphs[0].add_run(head).bold = True
    set_width(cell, [1.55, 4.95][i])
roles = [
    ("Student", "Submit and track service requests, book appointments, receive notifications, view FAQs, and use the chatbot."),
    ("Organization Representative", "Keeps normal Student access plus Event Request access after verified assignment by the Super Admin."),
    ("Admin Staff", "Processes operational requests, appointments, complaints, referrals, ID applications, and help desk tickets."),
    ("Super Admin", "Manages staff accounts, masterlist, organizations, representatives, settings, reports, backups, FAQs, and email blast."),
    ("Scanner Personnel", "Verifies secure QR codes and completes document claims without broad administrative access."),
]
for role, access in roles:
    cells = table.add_row().cells
    cells[0].text = role
    cells[1].text = access
    cells[0].paragraphs[0].runs[0].bold = True
    set_width(cells[0], 1.55); set_width(cells[1], 4.95)

add_callout("Important defense point:", "Organization Representative is an additional permission, not a replacement for the Student role. Only active representatives can view and submit Event Requests.")

doc.add_heading("4. Technology Stack", level=1)
table = doc.add_table(rows=1, cols=2)
clean_table(table)
table.style = "Table Grid"
for i, head in enumerate(["Technology", "Purpose in QR-SASMS"]):
    cell = table.rows[0].cells[i]
    shade(cell, LIGHT_BLUE)
    cell.paragraphs[0].add_run(head).bold = True
    set_width(cell, [1.8, 4.7][i])
stack = [
    ("Next.js 14", "Full-stack web framework for pages, React interface, and server API routes."),
    ("React", "Builds dynamic and reusable user interface components."),
    ("TypeScript / JavaScript", "Implements application logic and type-safe development."),
    ("Node.js", "Runs server-side application logic through Next.js."),
    ("PostgreSQL", "Stores related system records such as users, requests, appointments, notifications, and logs."),
    ("Prisma ORM", "Connects application code to PostgreSQL and manages schema-based database access."),
    ("bcrypt", "Hashes passwords before storage."),
    ("JWT / jose", "Supports authentication and secure server-validated QR tokens."),
    ("QR libraries", "Generates QR codes for document claiming."),
    ("CSS", "Provides the custom responsive QR-SASMS interface."),
    ("GitHub", "Version control and collaboration."),
    ("Render", "Hosts the frontend and backend web application."),
    ("SMTP provider", "Sends reset-password, request-status, and reminder emails after configuration."),
]
for tech, purpose in stack:
    cells = table.add_row().cells
    cells[0].text = tech
    cells[1].text = purpose
    cells[0].paragraphs[0].runs[0].bold = True
    set_width(cells[0], 1.8); set_width(cells[1], 4.7)

doc.add_heading("5. QR Verification and Claiming", level=1)
add_label_paragraph("Why QR? ", "The QR code is used for secure document claiming. It is generated only when a request becomes Ready to Claim.")
doc.add_paragraph("Key security controls:")
for item in [
    "The QR contains a unique secure token rather than plain student or document information.",
    "The server validates the token before a document can be released.",
    "The token can expire and is intended for one-time claiming.",
    "A successful claim marks the request Completed and creates an audit-log record.",
    "Scanner Personnel have limited access focused on verification and completion.",
]: add_bullet(item)

doc.add_heading("6. Appointment Management", level=1)
doc.add_paragraph("The system prevents duplicate appointments by checking schedule conditions before an appointment is saved.")
for item in [
    "Booked or full slots are disabled or crossed out.",
    "Capacity can be configured per time slot.",
    "A student can be restricted from making more than one appointment on the same date.",
    "Business hours, holidays, booking cutoffs, rescheduling, and cancellation rules can be managed through system settings.",
]: add_bullet(item)

doc.add_heading("7. Event Request Authorization", level=1)
doc.add_paragraph("Not all students can submit an Event Request. This protects official student-organization transactions from unauthorized submissions.")
for item in [
    "Super Admin registers or activates the student organization.",
    "Super Admin assigns verified student officers as Organization Representatives.",
    "Only active representatives can see and submit Event Requests.",
    "Representative access can be revoked, removed, reassigned, or renewed per academic year.",
    "The request is associated with the organization, representative, adviser information, and audit trail.",
]: add_number(item)

doc.add_heading("8. Chatbot and Help Desk", level=1)
doc.add_paragraph("The chatbot provides clickable FAQ answers for common questions. When the concern is unique or cannot be resolved by FAQ content, it can be escalated into a trackable Help Desk ticket for staff response.")
add_callout("Benefit:", "A student concern is not simply ignored. It can move from FAQ guidance to a staff-managed ticket with a visible status and response history.")

doc.add_heading("9. Database, Security, and Audit Controls", level=1)
doc.add_heading("Password security", level=2)
doc.add_paragraph("Passwords are hashed with bcrypt before storage. The system must never save passwords as plain text.")
doc.add_heading("Authorization", level=2)
doc.add_paragraph("Server routes enforce role-based access so Students, Admin Staff, Super Admins, and Scanner Personnel only reach functions allowed for their role.")
doc.add_heading("File access", level=2)
doc.add_paragraph("Students should only access their own submitted files and request records. Staff access should be limited to files required for processing.")
doc.add_heading("Audit logs", level=2)
doc.add_paragraph("Audit logs create accountability by recording major actions such as approval, rejection, QR claim completion, organization assignment, staff-management actions, and important configuration changes.")

doc.add_heading("10. Deployment and Email Explanation", level=1)
add_label_paragraph("Deployment: ", "GitHub stores the source code, while Render deploys and hosts the web application. PostgreSQL stores persistent data.")
add_label_paragraph("Email: ", "Render hosts the application but is not an email-sending service. Real emails require an external SMTP provider configured through secure environment variables. SMTP credentials must never be committed to GitHub.")
add_callout("Honest limitation:", "Production email delivery depends on correct SMTP configuration and provider limitations. A deployment platform does not automatically provide unlimited email delivery.")

doc.add_heading("11. Likely Panel Questions and Suggested Answers", level=1)
qas = [
    ("What makes QR-SASMS different from a normal request form?", "It combines request management, appointments, role-based access, chatbot-to-ticket escalation, organization representative authorization, audit logging, and secure QR claiming in one platform."),
    ("Why use QR codes for document claiming?", "QR verification makes claiming faster and traceable. The server validates a one-time secure token before the document is released, which reduces unauthorized or repeated claims."),
    ("How does the system prevent double booking?", "The system checks the date, time slot, configured capacity, and student booking rules before saving an appointment. Unavailable slots are disabled."),
    ("Can every student submit an Event Request?", "No. Only students assigned as active Organization Representatives by the Super Admin can access Event Requests."),
    ("What happens if a request is rejected?", "The admin provides a rejection reason. The student sees it, corrects the requirements, and can re-upload or resubmit for review."),
    ("Why use PostgreSQL?", "The system has highly related records such as users, requests, appointments, organizations, notifications, and logs. PostgreSQL supports relational integrity and reliable structured storage."),
    ("What is Prisma ORM?", "Prisma is an ORM that lets the application work with PostgreSQL through structured TypeScript or JavaScript code while keeping the database schema organized."),
    ("What is the purpose of audit logs?", "They record who performed important actions and when. This supports accountability, troubleshooting, and monitoring."),
    ("What happens when the chatbot cannot answer?", "The concern can be escalated into a Help Desk ticket so an authorized staff member can respond and update its status."),
    ("What are the current production requirements?", "The system needs a secure production database, configured environment variables, an SMTP provider for real email, controlled user testing, and role and authorization testing before full deployment."),
]
for q, a in qas: add_qa(q, a)

doc.add_heading("12. Short Developer Introduction", level=1)
doc.add_paragraph(
    "As the developer, my responsibility is to design and implement the user interface, backend API routes, database integration, authentication, role-based authorization, QR verification workflow, appointment validation, notifications, and audit logging. My goal is to ensure that these modules work together as one secure and organized system for Student Services."
)

doc.add_heading("13. Suggested Closing Statement", level=1)
add_callout("Closing:", "QR-SASMS is not only a digital request form. It is an integrated student-service management platform that provides transparent request tracking for students and a more organized, secure, and traceable workflow for the Student Services Office.")

doc.core_properties.title = "QR-SASMS Developer Defense Reviewer"
doc.core_properties.subject = "Capstone research defense reviewer"
doc.core_properties.author = "QR-SASMS Capstone Team"
doc.save(OUT)
print(OUT)
